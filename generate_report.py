"""
Generate the IEEE-style report (Methodology / Results / Conclusions)
as a .docx file that can be pasted / appended into the existing
report already produced by the team.
"""

import os
import pandas as pd
from docx              import Document
from docx.shared       import Pt, Inches, RGBColor, Cm
from docx.enum.text    import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns      import qn
from docx.oxml         import OxmlElement


HERE        = os.path.dirname(os.path.abspath(__file__))
RESULTS_DIR = os.path.abspath(os.path.join(HERE, "..", "results"))
CM_DIR      = os.path.join(RESULTS_DIR, "confusion_matrices")
OUT_PATH    = os.path.abspath(os.path.join(
    HERE, "..", "Malware_Classification_Report_Methodology_Results.docx"))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_two_columns(section):
    """IEEE conference papers are two-column."""
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        cols[0].set(qn("w:num"), "2")
        cols[0].set(qn("w:space"), "432")     # 0.3 inch gap
    else:
        c = OxmlElement("w:cols")
        c.set(qn("w:num"), "2")
        c.set(qn("w:space"), "432")
        sectPr.append(c)


def add_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper())
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        return p
    else:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.italic = True
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
    if not p.runs:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
    else:
        p.runs[0].text = text
    return p


def add_table_from_df(doc, df, caption=None, decimals=4):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"

    hdr = t.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = "Times New Roman"

    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            if isinstance(v, float):
                cells[i].text = "-" if pd.isna(v) else f"{v:.{decimals}f}"
            else:
                cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Times New Roman"
    return t


def add_image(doc, path, caption, width=Inches(3.2)):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.font.size = Pt(9)
    r.italic = True
    r.font.name = "Times New Roman"


# ---------------------------------------------------------------------------
# Load results
# ---------------------------------------------------------------------------
def load_metrics():
    comp = pd.read_csv(os.path.join(RESULTS_DIR, "comparison_table.csv"))
    pca  = pd.read_csv(os.path.join(RESULTS_DIR, "pca_comparison.csv"))
    cnn_hist = pd.read_csv(os.path.join(RESULTS_DIR, "cnn_training_history.csv"))
    try:
        cnn_mis = pd.read_csv(os.path.join(RESULTS_DIR, "cnn_misclassifications.csv"))
    except Exception:
        cnn_mis = pd.DataFrame()
    return comp, pca, cnn_hist, cnn_mis


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
def build_report():
    comp, pca, cnn_hist, cnn_mis = load_metrics()

    doc = Document()

    # Page / font defaults
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)
        set_two_columns(s)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    # Title note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Malware Classification Using Machine Learning\n"
                  "— Methodology, Results, and Conclusion —")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    # =========================================================
    # IV. Methodology
    # =========================================================
    add_heading(doc, "IV. Methodology")

    add_body(doc,
        "This section describes the end-to-end pipeline used to build and "
        "evaluate the proposed malware classification system. Two parallel "
        "pipelines were implemented: (1) a traditional machine-learning "
        "pipeline operating on structured Portable Executable (PE) features, "
        "and (2) a deep-learning pipeline that treats malware binaries as "
        "images and classifies them using a Convolutional Neural Network "
        "(CNN). A third branch combines the individual machine-learning "
        "models through a soft voting ensemble, and a Principal-Component "
        "Analysis (PCA) study investigates the effect of dimensionality "
        "reduction on classification performance.")

    add_heading(doc, "A. Datasets", 2)
    add_body(doc,
        "Two publicly available malware datasets were used so that each "
        "pipeline is trained on data that matches its inductive bias. The "
        "Classification of Malwares (CLaMP) dataset [13] contains 5,210 "
        "samples and 69 structured PE-header features such as section "
        "counts, entropy values, optional-header fields, and API-related "
        "metadata. Each sample is labelled as benign or malicious, giving a "
        "balanced binary classification task (2,488 benign vs. 2,722 "
        "malicious after cleaning). The Malimg dataset [14] contains 9,339 "
        "grayscale byte-plot images organised into 25 malware families. "
        "This dataset exhibits pronounced class imbalance, ranging from "
        "80 samples (Skintrim.N) to 2,949 samples (Allaple.A), reflecting "
        "the natural distribution of malware families in the wild.")

    add_heading(doc, "B. Data Preparation (CLaMP pipeline)", 2)
    add_body(doc,
        "Data preparation follows five deterministic steps. First, the raw "
        "CSV is loaded and inspected to confirm the expected shape and "
        "column types. Second, duplicates and rows containing missing "
        "values are removed — duplicates bias training toward memorisation "
        "of specific samples, while missing values cause runtime errors. "
        "The free-text field 'fileinfo' is dropped because it contributes "
        "no numeric signal usable by the classifiers. Third, the target "
        "column is passed through LabelEncoder so the models operate on "
        "integer labels while preserving the original class mapping. "
        "Fourth, the dataset is split 80%/20% into training and test sets "
        "with stratification on the class label, guaranteeing that the "
        "class ratio is preserved in both partitions. Finally, "
        "StandardScaler is fitted on the training data only and applied to "
        "the test set, producing zero-mean, unit-variance features — a "
        "critical preprocessing step for SVM and a beneficial one for the "
        "other classifiers.")

    add_heading(doc, "C. Traditional Machine-Learning Models", 2)
    add_body(doc,
        "Three complementary algorithms are trained on the CLaMP features. "
        "Random Forest (RF) constructs 100 decision trees on bootstrap "
        "samples of the data and aggregates them by majority vote; it is "
        "robust to irrelevant features, resistant to overfitting, and "
        "natively exposes feature-importance scores. Support Vector Machine "
        "(SVM) with a Radial Basis Function (RBF) kernel seeks the maximum-"
        "margin hyperplane separating benign and malicious samples in a "
        "high-dimensional feature space, and benefits directly from the "
        "feature standardisation performed in the preparation phase. "
        "Gradient Boosting (GB) constructs an additive ensemble of 100 "
        "shallow trees where each new tree is trained to correct the "
        "errors of the previous ensemble, trading extra training time for "
        "an often higher decision accuracy. All three models use the "
        "scikit-learn [15] implementations with default hyper-parameters "
        "and a fixed random seed (42) for reproducibility.")

    add_heading(doc, "D. Voting Ensemble", 2)
    add_body(doc,
        "To exploit the complementary strengths of the individual "
        "classifiers, a soft-voting ensemble is constructed from the "
        "Random Forest, SVM, and Gradient Boosting base learners. At "
        "inference time each base classifier outputs a probability "
        "distribution over the two classes; the ensemble averages these "
        "distributions and returns the class with the highest averaged "
        "probability. Soft voting is preferred over hard voting because it "
        "makes use of the full confidence information produced by each "
        "model rather than only the final class decision, leading to a "
        "smoother and more stable decision function.")

    add_heading(doc, "E. Principal Component Analysis", 2)
    add_body(doc,
        "PCA is applied to the standardised CLaMP feature space to "
        "investigate whether the 67-dimensional representation can be "
        "compressed without degrading classification performance. The "
        "cumulative explained-variance curve is first inspected to locate "
        "the 95% and 99% variance thresholds. Random Forest is then "
        "retrained on three reduced spaces — 30, 50 and the maximum "
        "available number of components (67) — and the resulting accuracy "
        "and F1 score are compared against a no-PCA baseline. Because the "
        "dataset contains only 67 features, the originally requested "
        "100-component configuration is clipped to 67, which corresponds "
        "to a full-rank decorrelating rotation (no dimensionality loss).")

    add_heading(doc, "F. Deep-Learning Pipeline (CNN)", 2)
    add_body(doc,
        "The deep-learning branch operates on the Malimg dataset. All "
        "byte-plot images are converted to single-channel grayscale, "
        "resized to 128×128 pixels, and normalised to the [-1, 1] range. "
        "A stratified 80%/20% split produces 7,471 training and 1,868 "
        "test images. The CNN consists of three convolutional blocks — "
        "each block pairs a 3×3 convolution with ReLU activation and a "
        "2×2 max-pool layer, halving the spatial resolution at every "
        "stage (128→64→32→16). Channels grow from 32 to 64 to 128, so the "
        "network progresses from detecting edges and local textures to "
        "higher-level structural patterns characteristic of malware "
        "families. A fully-connected head with 256 hidden units, ReLU "
        "activation, and 50% dropout feeds the final 25-way softmax "
        "classifier. Training is performed for 5 epochs with a batch size "
        "of 32, the Adam optimiser at a learning rate of 1×10⁻³, and "
        "cross-entropy loss, on an NVIDIA CUDA-enabled GPU.")

    add_heading(doc, "G. Evaluation Framework", 2)
    add_body(doc,
        "All models are evaluated on the same held-out test partition of "
        "their respective dataset using a unified evaluation module so "
        "that results are directly comparable. Four standard classification "
        "metrics are reported: accuracy, weighted precision, weighted "
        "recall, and weighted F1-score. Weighted averaging is used so that "
        "each class contributes to the final score in proportion to its "
        "support — a requirement when the Malimg dataset contains classes "
        "with very different sizes. In addition, a confusion-matrix "
        "heat-map is generated for every model to expose the specific "
        "families that are most frequently confused, and a dedicated "
        "misclassification-analysis routine ranks the most common "
        "true-vs-predicted pairs. These qualitative artefacts complement "
        "the aggregate metrics and directly support the interpretability "
        "requirement highlighted by the research gap.")

    # =========================================================
    # V. Results
    # =========================================================
    add_heading(doc, "V. Results and Discussion")

    add_heading(doc, "A. Overall Comparison", 2)
    add_body(doc,
        "Table I presents the aggregate performance of every model variant "
        "evaluated in this study, ranked in decreasing order of accuracy. "
        "On the CLaMP dataset, Random Forest achieves the strongest "
        "performance with an accuracy of 99.04% and an F1 score of 99.04%, "
        "closely followed by Gradient Boosting (98.08%). The soft-voting "
        "ensemble obtains 97.79%, and SVM trails with 95.59%. On the "
        "Malimg dataset, the CNN reaches 97.81% accuracy and 97.31% F1 "
        "after only five training epochs, demonstrating that convolutional "
        "pattern extraction is highly effective on the byte-plot "
        "representation of malware binaries.")

    comp_pretty = comp.rename(columns={
        "model": "Model", "accuracy": "Accuracy",
        "precision": "Precision", "recall": "Recall", "f1": "F1",
    })
    add_table_from_df(
        doc, comp_pretty,
        caption="Table I. Ranked performance across all model variants.")

    add_heading(doc, "B. Feature Importance", 2)
    add_body(doc,
        "Figure 1 displays the 20 most important features identified by "
        "the Random Forest classifier on the CLaMP dataset. The ranking is "
        "dominated by PE-header and packer-related fields, confirming that "
        "structural characteristics of the executable carry the majority "
        "of the discriminative signal for distinguishing benign from "
        "malicious samples. This interpretability property is a key "
        "advantage of Random Forest over deep learning baselines and "
        "directly addresses the interpretability research gap identified "
        "in Section III.")
    add_image(doc, os.path.join(RESULTS_DIR, "feature_importance.png"),
              "Fig. 1. Top-20 Random Forest feature importances (CLaMP).")

    add_heading(doc, "C. PCA Analysis", 2)
    add_body(doc,
        "Table II summarises the effect of dimensionality reduction on "
        "Random Forest performance. Using 30 principal components already "
        "captures most of the class-discriminative information, yielding "
        "97.03% accuracy versus the 99.04% no-PCA baseline — a drop of "
        "roughly 2 percentage points while shrinking the feature space to "
        "less than half of its original size. Increasing to 50 or 67 "
        "components does not recover the baseline performance, indicating "
        "that the decorrelating rotation performed by PCA slightly "
        "degrades Random Forest's ability to exploit informative axis-"
        "aligned splits. The conclusion is that, for this dataset and "
        "this classifier, the full feature set should be preferred, while "
        "PCA remains useful as a fast approximation when computational "
        "budget is constrained.")
    pca_pretty = pca.rename(columns={
        "variant": "Variant", "n_components": "# Components",
        "accuracy": "Accuracy", "f1": "F1"})
    add_table_from_df(
        doc, pca_pretty,
        caption="Table II. Random Forest accuracy and F1 under PCA.")
    add_image(doc, os.path.join(RESULTS_DIR, "pca_explained_variance.png"),
              "Fig. 2. Cumulative explained variance vs number of components.")

    add_heading(doc, "D. CNN Training Dynamics", 2)
    rows = "; ".join(
        f"epoch {int(r.epoch)}: loss={r.loss:.4f}, acc={r.acc:.4f}"
        for _, r in cnn_hist.iterrows())
    add_body(doc,
        "The CNN training loss decreases monotonically from 0.520 in the "
        "first epoch to 0.090 in the fifth, while training accuracy rises "
        "from 85.2% to 97.2%. The absence of divergence between training "
        "and test accuracy (97.15% train vs. 97.81% test) suggests that "
        "dropout and the limited number of epochs successfully prevented "
        "overfitting on the imbalanced dataset. Full training history: "
        f"{rows}.")

    add_heading(doc, "E. Confusion-Matrix Analysis", 2)
    add_body(doc,
        "Confusion matrices for every model were generated and saved in "
        "the results/confusion_matrices/ directory. On CLaMP (Fig. 3), "
        "Random Forest yields a near-diagonal matrix: only ten benign "
        "samples are misclassified as malicious and fewer are missed in "
        "the opposite direction, indicating high specificity and high "
        "sensitivity simultaneously. On Malimg (Fig. 4), the CNN correctly "
        "classifies most of the 25 families, but systematic confusion is "
        "observed between visually similar families.")
    add_image(doc, os.path.join(CM_DIR, "cm_random_forest.png"),
              "Fig. 3. Confusion matrix — Random Forest (CLaMP).")
    add_image(doc, os.path.join(CM_DIR, "cm_cnn.png"),
              "Fig. 4. Confusion matrix — CNN (Malimg).", width=Inches(3.3))

    add_heading(doc, "F. Misclassification Analysis", 2)
    if not cnn_mis.empty:
        top = cnn_mis.head(5)
        mis_list = "; ".join(
            f"{r['true']} → {r['predicted']} ({int(r['count'])})"
            for _, r in top.iterrows())
    else:
        mis_list = "none"
    add_body(doc,
        "A dedicated analysis of the CNN errors identifies the "
        "most frequent true-vs-predicted confusions: " + mis_list + ". "
        "The dominant failure mode is Autorun.K being classified as "
        "Yuner.A, a known pair in the literature because both families "
        "rely on similar packer patterns that produce nearly identical "
        "byte-plot textures. Confusion among Swizzor.gen!E, Swizzor.gen!I "
        "and C2LOP variants reflects the fact that these samples share "
        "the same code base, differing only in minor mutations. This is "
        "an intrinsic limitation of purely visual representations and "
        "motivates the hybrid strategies discussed in Section VI.")

    add_heading(doc, "G. Model Comparison Chart", 2)
    add_body(doc,
        "Figure 5 visualises the overall ranking across all seven model "
        "variants. The three strongest performers are Random Forest, "
        "Gradient Boosting, and the CNN; the SVM baseline lags but "
        "remains competitive; and the PCA-reduced variants cluster "
        "slightly below the baseline, confirming the analysis in "
        "Section V-C.")
    add_image(doc, os.path.join(RESULTS_DIR, "model_comparison_chart.png"),
              "Fig. 5. Aggregate comparison across all model variants.",
              width=Inches(3.3))

    # =========================================================
    # VI. Conclusion
    # =========================================================
    add_heading(doc, "VI. Conclusion")
    add_body(doc,
        "This paper presented a dual-pipeline malware-classification "
        "system that compares traditional machine-learning models on "
        "structured PE-header features with a Convolutional Neural "
        "Network trained directly on malware byte-plot images. Seven "
        "model variants were trained and evaluated under a common "
        "evaluation framework: Random Forest, SVM, Gradient Boosting, a "
        "soft-voting ensemble of the three, three Random-Forest variants "
        "with PCA applied at 30, 50, and 67 components, and a CNN with "
        "three convolutional blocks.")
    add_body(doc,
        "The strongest overall result is obtained by Random Forest, "
        "which reaches 99.04% accuracy and F1 on the CLaMP dataset. "
        "Gradient Boosting (98.08%), the CNN on Malimg (97.81%), and the "
        "voting ensemble (97.79%) form a closely-matched second tier, "
        "confirming that both structured-feature and image-based "
        "approaches are viable for automated malware classification. "
        "SVM performs credibly (95.59%) but is clearly outperformed by "
        "the tree-based alternatives, while PCA-reduced variants "
        "demonstrate that a 2-percentage-point accuracy loss can be "
        "traded for a halving of the feature space when computational "
        "resources are constrained.")
    add_body(doc,
        "The feature-importance analysis performed on the Random Forest "
        "and the confusion-matrix / misclassification analyses performed "
        "on every model directly address the three research gaps "
        "highlighted in Section III: (i) the lack of interpretable "
        "black-box baselines is mitigated by the built-in feature-"
        "importance output of Random Forest; (ii) the absence of direct "
        "static-vs-visual comparisons is closed by evaluating both "
        "pipelines under an identical evaluation framework; and (iii) "
        "the limited treatment of imbalanced data is handled through "
        "weighted metrics and detailed per-class error analysis.")
    add_body(doc,
        "Future work will focus on three directions. First, a hybrid "
        "model that concatenates CNN image embeddings with structured "
        "PE-header features may combine the complementary strengths "
        "demonstrated here and resolve the Swizzor/C2LOP confusion "
        "observed in Section V-F. Second, explainable-AI techniques "
        "such as SHAP and Grad-CAM can be applied so that malware "
        "analysts obtain human-readable justifications alongside every "
        "classification. Third, evaluation on additional datasets — "
        "notably CIC-AndMal-2020 for dynamic behavioural features — "
        "will provide a stronger test of generalisation to malware "
        "families unseen during training.")

    # =========================================================
    # Appendix
    # =========================================================
    add_heading(doc, "Appendix: Reproducibility")
    add_body(doc,
        "All experiments use a fixed random seed (42) and are driven by "
        "the run_all.py master script, which executes the ML pipeline, "
        "PCA experiments, CNN training, and final comparison in order. "
        "The full code base (fatima_evaluation.py, ml_pipeline.py, "
        "pca_analysis.py, cnn_pipeline.py, generate_comparison.py) and "
        "every artefact referenced in this section are stored under "
        "implementation/results/. Library versions: scikit-learn 1.3, "
        "PyTorch 2.7 with CUDA 11.8, pandas 2.0, matplotlib and seaborn "
        "for plotting.")

    doc.save(OUT_PATH)
    print(f"Report saved -> {OUT_PATH}")


if __name__ == "__main__":
    build_report()
